# Testes do módulo de extração (app/extracao/extrator.py).
# Os arquivos de exemplo (PDF e DOCX) são gerados aqui mesmo, num diretório
# temporário do pytest (tmp_path): não precisamos guardar arquivos binários
# de exemplo no repositório.

import fitz
import pytest
from docx import Document

from app.extracao.extrator import extrair_texto


def criar_pdf_exemplo(caminho) -> None:
    """Cria um PDF de duas páginas com texto conhecido em cada uma."""
    documento = fitz.open()
    pagina1 = documento.new_page()
    pagina1.insert_text((72, 72), "Objeto do contrato: aquisicao de bens.")
    pagina2 = documento.new_page()
    pagina2.insert_text((72, 72), "Clausula segunda: prazo de execucao.")
    documento.save(str(caminho))
    documento.close()


def criar_docx_exemplo(caminho) -> None:
    """Cria um DOCX com três parágrafos, em ordem conhecida."""
    documento = Document()
    documento.add_paragraph("Primeiro paragrafo do edital.")
    documento.add_paragraph("Segundo paragrafo, com a clausula principal.")
    documento.add_paragraph("Terceiro e ultimo paragrafo.")
    documento.save(str(caminho))


def test_pdf_retorna_numero_de_paginas_e_texto_no_lugar_certo(tmp_path):
    caminho_pdf = tmp_path / "exemplo.pdf"
    criar_pdf_exemplo(caminho_pdf)

    resultado = extrair_texto(str(caminho_pdf))

    assert resultado.tipo == "pdf"
    assert resultado.num_paginas == 2
    assert len(resultado.blocos) == 2

    # O texto conhecido da página 2 deve aparecer no bloco da página 2.
    bloco_pagina_2 = resultado.blocos[1]
    assert bloco_pagina_2.pagina == 2
    assert bloco_pagina_2.localizador == "página 2"
    assert "Clausula segunda" in bloco_pagina_2.texto


def test_docx_retorna_blocos_com_pagina_nula(tmp_path):
    caminho_docx = tmp_path / "exemplo.docx"
    criar_docx_exemplo(caminho_docx)

    resultado = extrair_texto(str(caminho_docx))

    assert resultado.tipo == "docx"
    assert resultado.num_paginas is None
    assert len(resultado.blocos) == 3

    for bloco in resultado.blocos:
        assert bloco.pagina is None

    assert resultado.blocos[1].localizador == "parágrafo 2"
    assert "clausula principal" in resultado.blocos[1].texto


def test_extensao_nao_suportada_gera_erro_claro(tmp_path):
    caminho_invalido = tmp_path / "edital.xyz"
    caminho_invalido.write_text("conteudo qualquer")

    with pytest.raises(ValueError, match=r"formato não suportado: \.xyz"):
        extrair_texto(str(caminho_invalido))
