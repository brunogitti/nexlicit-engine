# Testes da geração da minuta de proposta (Fase 4, Camada 1, 19/08/2026,
# app/geracao/minuta.py). Testes de unidade puros (dicts montados na mão,
# sem banco -- mesmo padrão de tests/test_declaracoes.py) mais um golden
# test contra o catálogo real de Paulínia (offline, sem IA, mesmo padrão
# de tests/test_planilha_preco.py).

import io
from pathlib import Path

import pytest

from app.extracao.extrator import extrair_texto
from app.extracao.tabela_itens import localizar_documento_com_tabela
from app.geracao.assinatura import EmpresaSemRepresentanteError
from app.geracao.minuta import gerar_minuta
from app.geracao.planilha_preco import SemCatalogoError

PASTA_UPLOADS_REAIS = Path(__file__).resolve().parent.parent / "uploads"
CAMINHO_PAULINIA = PASTA_UPLOADS_REAIS / "7" / "PE 82.2026 - PM PAULINIA - 15.06 - MAXXIMED - BNC - 18974758 - SQ - SIM OK IMPRESSO.pdf"


def _processo(**sobrescreve) -> dict:
    dados = {"id": 1, "nome": "Pregão Eletrônico 01/2026", "orgao": None, "modalidade": None, "objeto": None, "validade_proposta": None}
    dados.update(sobrescreve)
    return dados


def _empresa(**sobrescreve) -> dict:
    dados = {
        "id": 1,
        "razao_social": "Exemplo Fornecedora de Materiais Ltda",
        "cnpj": "12.345.678/0001-90",
        "representante_legal_nome": "José da Silva Fictício",
        "representante_legal_cargo": "Sócio-administrador",
        "representante_legal_cpf": "000.000.000-00",
    }
    dados.update(sobrescreve)
    return dados


def _catalogo(*numeros_e_textos: tuple[int, str]) -> list[dict]:
    return [{"numero": n, "texto_bruto": t, "pagina": 1, "localizador": "página 1"} for n, t in numeros_e_textos]


def _textos(documento) -> list[str]:
    return [p.text for p in documento.paragraphs]


def _linhas_tabela(documento) -> list[list[str]]:
    tabela = documento.tables[0]
    return [[celula.text for celula in linha.cells] for linha in tabela.rows]


def test_sem_catalogo_levanta_erro():
    with pytest.raises(SemCatalogoError, match="1"):
        gerar_minuta(_processo(), _empresa(), [], {})


def test_empresa_sem_representante_levanta_erro():
    catalogo = _catalogo((1, "Item qualquer"))
    with pytest.raises(EmpresaSemRepresentanteError):
        gerar_minuta(_processo(), _empresa(representante_legal_nome=None), catalogo, {})


def test_cabecalho_tem_dados_da_empresa_e_do_processo():
    catalogo = _catalogo((1, "Item qualquer"))
    processo = _processo(orgao="Prefeitura X", modalidade="Pregão Eletrônico", objeto="Aquisição de materiais")

    documento = gerar_minuta(processo, _empresa(), catalogo, {})

    textos = " ".join(_textos(documento))
    assert "Exemplo Fornecedora de Materiais Ltda" in textos
    assert "12.345.678/0001-90" in textos
    assert "Pregão Eletrônico 01/2026" in textos
    assert "Prefeitura X" in textos
    assert "Pregão Eletrônico" in textos
    assert "Aquisição de materiais" in textos


def test_tabela_tem_uma_linha_por_item_mais_cabecalho():
    catalogo = _catalogo((1, "Cadeira de rodas"), (2, "Muleta axilar"))
    precos = {
        1: {"quantidade": 5, "preco_unitario": 200.0, "marca": "MarcaX", "fabricante": None, "modelo": "M1"},
        2: {"quantidade": 20, "preco_unitario": 15.0, "marca": None, "fabricante": None, "modelo": None},
    }

    documento = gerar_minuta(_processo(), _empresa(), catalogo, precos)
    linhas = _linhas_tabela(documento)

    assert len(linhas) == 3  # cabeçalho + 2 itens
    assert linhas[1][0] == "1"
    assert linhas[1][1] == "Cadeira de rodas"
    assert "Marca: MarcaX" in linhas[1][2]
    assert "Modelo: M1" in linhas[1][2]
    assert linhas[1][3] == "5"
    assert linhas[1][4] == "200"
    assert linhas[1][5] == "1000"  # 5 * 200


def test_item_sem_marca_fabricante_modelo_fica_com_celula_vazia():
    catalogo = _catalogo((1, "Item sem dado de produto"))
    precos = {1: {"quantidade": 1, "preco_unitario": 10.0, "marca": None, "fabricante": None, "modelo": None}}

    documento = gerar_minuta(_processo(), _empresa(), catalogo, precos)
    linhas = _linhas_tabela(documento)

    assert linhas[1][2] == ""  # célula de marca/fabricante/modelo vazia, nada inventado


def test_item_sem_preco_fica_com_celulas_vazias_sem_travar_a_minuta():
    catalogo = _catalogo((1, "Item sem preço"), (2, "Item com preço"))
    precos = {2: {"quantidade": 3, "preco_unitario": 10.0, "marca": None, "fabricante": None, "modelo": None}}

    documento = gerar_minuta(_processo(), _empresa(), catalogo, precos)
    linhas = _linhas_tabela(documento)

    assert linhas[1][3] == "" and linhas[1][4] == "" and linhas[1][5] == ""  # item 1: nada inventado
    assert linhas[2][3] == "3" and linhas[2][4] == "10" and linhas[2][5] == "30"  # item 2: calculado normalmente


def test_item_com_so_quantidade_nao_calcula_total_parcial():
    catalogo = _catalogo((1, "Item parcial"))
    precos = {1: {"quantidade": 10, "preco_unitario": None, "marca": None, "fabricante": None, "modelo": None}}

    documento = gerar_minuta(_processo(), _empresa(), catalogo, precos)
    linhas = _linhas_tabela(documento)

    assert linhas[1][3] == "10"
    assert linhas[1][4] == ""
    assert linhas[1][5] == ""  # sem os dois números, não inventa metade da conta


def test_validade_preenchida_aparece_na_clausula():
    catalogo = _catalogo((1, "Item qualquer"))
    documento = gerar_minuta(_processo(validade_proposta="60 dias"), _empresa(), catalogo, {})

    textos = " ".join(_textos(documento))
    assert "60 dias" in textos
    assert "PENDENTE" not in textos


def test_validade_ausente_gera_marca_visivel_de_pendencia():
    catalogo = _catalogo((1, "Item qualquer"))
    documento = gerar_minuta(_processo(validade_proposta=None), _empresa(), catalogo, {})

    textos = " ".join(_textos(documento))
    assert "PENDENTE" in textos
    assert "validade da proposta" in textos.lower()


def test_bloco_de_assinatura_tem_nome_cargo_e_cpf():
    catalogo = _catalogo((1, "Item qualquer"))
    documento = gerar_minuta(_processo(), _empresa(), catalogo, {})

    textos = _textos(documento)
    assert any("José da Silva Fictício, Sócio-administrador" in t for t in textos)
    assert any("CPF: 000.000.000-00" in t for t in textos)


def test_devolve_docx_reabrivel():
    catalogo = _catalogo((1, "Item qualquer"))
    documento = gerar_minuta(_processo(), _empresa(), catalogo, {})

    buffer = io.BytesIO()
    documento.save(buffer)
    buffer.seek(0)

    from docx import Document

    reaberto = Document(buffer)
    assert len(reaberto.tables) == 1


# ---------- Teste de ouro: minuta contra catálogo real de Paulínia ----------


@pytest.mark.skipif(not CAMINHO_PAULINIA.exists(), reason="PDF real de Paulínia não está disponível localmente")
def test_gerar_minuta_contra_catalogo_real_de_paulinia_com_dado_de_teste():
    documento_extraido = extrair_texto(str(CAMINHO_PAULINIA))
    resultado = localizar_documento_com_tabela([documento_extraido])
    assert resultado is not None
    _, itens = resultado

    catalogo = [
        {"numero": item.numero, "texto_bruto": item.texto, "pagina": item.pagina, "localizador": item.localizador}
        for item in itens
    ]
    # Preço/marca de teste só no primeiro item -- confirma que o resto
    # sai em branco, sem travar a geração, mesmo padrão do golden test da
    # planilha de preço.
    precos = {
        itens[0].numero: {
            "quantidade": 10, "preco_unitario": 99.9,
            "marca": "Marca de teste", "fabricante": None, "modelo": None,
        }
    }
    processo = _processo(id=999, nome="PM Paulínia (teste)", validade_proposta="60 dias")

    documento = gerar_minuta(processo, _empresa(), catalogo, precos)
    linhas = _linhas_tabela(documento)

    assert len(linhas) == len(catalogo) + 1  # cabeçalho + itens
    assert linhas[1][3] == "10"
    assert linhas[1][4] == "99.9"
    assert linhas[1][5] == "999"
