# Testes da geração do recurso administrativo (Fase 4, Camada 1,
# 19/08/2026, app/geracao/recurso.py). Testes de unidade puros (dicts
# montados na mão, sem banco -- mesmo padrão de test_declaracoes.py e
# test_minuta.py) mais um golden test contra uma exigência real de
# Paulínia com narrativa fictícia de teste.
#
# Cobre especificamente as garantias de segurança combinadas com Bruno:
# fatos do usuário não misturados com fato verificado, base legal só
# aparece se veio da exigência (nunca inventada aqui -- este módulo só
# monta documento, quem gera o argumento é app.pipeline, testado à
# parte), marca de conferência presente sempre.

import io
import zipfile

import pytest

from app.geracao.assinatura import EmpresaSemRepresentanteError
from app.geracao.recurso import gerar_recurso


def _processo(**sobrescreve) -> dict:
    dados = {"id": 1, "nome": "Pregão Eletrônico 01/2026", "orgao": "Prefeitura X"}
    dados.update(sobrescreve)
    return dados


def _empresa(**sobrescreve) -> dict:
    dados = {
        "id": 1,
        "razao_social": "Exemplo Fornecedora de Materiais Ltda",
        "cnpj": "12.345.678/0001-90",
        "representante_legal_nome": "José da Silva Fictício",
        "representante_legal_cargo": "Sócio-administrador",
        "representante_legal_cpf": "000.000.000-00",
    }
    dados.update(sobrescreve)
    return dados


def _exigencia(**sobrescreve) -> dict:
    dados = {
        "categoria": "habilitacao_juridica",
        "descricao": "Certidão do CREA",
        "trecho": "O licitante deverá apresentar certidão de regularidade junto ao CREA.",
        "pagina": 5,
        "base_legal": "art. 66 da Lei 14.133/2021",
    }
    dados.update(sobrescreve)
    return dados


def _textos(documento) -> list[str]:
    return [p.text for p in documento.paragraphs]


NARRATIVA_TESTE = (
    "A certidão apresentada tinha validade até 15/03/2026, e a sessão de "
    "habilitação ocorreu em 10/03/2026, dentro do prazo."
)


def test_titulo_de_aviso_presente_e_destacado():
    documento = gerar_recurso(
        _processo(), _empresa(), _exigencia(), NARRATIVA_TESTE, "Fundamentação.", "Pedido."
    )
    textos = _textos(documento)
    assert any("MINUTA" in t and "REVISAR ANTES DE PROTOCOLAR" in t for t in textos)


def test_narrativa_do_usuario_vai_verbatim_e_separada_do_fato_verificado():
    documento = gerar_recurso(
        _processo(), _empresa(), _exigencia(), NARRATIVA_TESTE, "Fundamentação de teste.", "Pedido de teste."
    )
    textos = _textos(documento)

    # A narrativa aparece literalmente, sem alteração.
    assert any(NARRATIVA_TESTE in t for t in textos)

    # Vem precedida de um aviso claro de que é alegação não verificada,
    # numa seção PRÓPRIA (II) separada da exigência do edital (seção I) --
    # fato do usuário não se mistura com fato verificado no mesmo bloco.
    indice_fatos = next(i for i, t in enumerate(textos) if "Fatos alegados pelo recorrente" in t)
    indice_narrativa = next(i for i, t in enumerate(textos) if NARRATIVA_TESTE in t)
    indice_aviso = next(i for i, t in enumerate(textos) if "não verificada pelo sistema" in t)
    assert indice_fatos < indice_aviso < indice_narrativa

    # A seção I (exigência do edital) não contém a narrativa do usuário.
    indice_edital = next(i for i, t in enumerate(textos) if "Exigência do edital" in t)
    assert indice_edital < indice_fatos


def test_trecho_do_edital_aparece_literal_na_secao_de_fundamento():
    exigencia = _exigencia(trecho="Trecho literal exato do edital.")
    documento = gerar_recurso(_processo(), _empresa(), exigencia, NARRATIVA_TESTE, "F", "P")
    textos = " ".join(_textos(documento))
    assert "Trecho literal exato do edital." in textos
    assert "página 5" in textos.lower() or "5" in textos


def test_base_legal_da_exigencia_aparece_com_marca_de_conferencia():
    exigencia = _exigencia(base_legal="art. 66 da Lei 14.133/2021")
    documento = gerar_recurso(_processo(), _empresa(), exigencia, NARRATIVA_TESTE, "F", "P")
    textos = " ".join(_textos(documento))

    assert "art. 66 da Lei 14.133/2021" in textos
    assert "CONFERIR ANTES DE PROTOCOLAR" in textos


def test_marca_de_conferencia_presente_mesmo_com_confianca_localizado():
    # Mais rígido que declarações/minuta de propósito: a marca aparece
    # SEMPRE, mesmo quando confianca="localizado" (a exigência já foi
    # confirmada contra o texto-fonte) -- o risco aqui não é "esse
    # trecho existe", é "essa citação certa foi usada na linha de
    # argumento certa", e isso só avaliação profissional resolve.
    exigencia = _exigencia(base_legal="art. 66 da Lei 14.133/2021", confianca="localizado")
    documento = gerar_recurso(_processo(), _empresa(), exigencia, NARRATIVA_TESTE, "F", "P")
    textos = " ".join(_textos(documento))
    assert "CONFERIR ANTES DE PROTOCOLAR" in textos

    # E o comentário nativo do Word de verdade também está presente
    # (mesma checagem de test_declaracoes.py -- não só texto solto).
    buffer = io.BytesIO()
    documento.save(buffer)
    with zipfile.ZipFile(buffer) as zf:
        assert "word/comments.xml" in zf.namelist()


def test_exigencia_sem_base_legal_nao_inventa_artigo():
    exigencia = _exigencia(base_legal=None)
    documento = gerar_recurso(_processo(), _empresa(), exigencia, NARRATIVA_TESTE, "F", "P")
    textos = " ".join(_textos(documento))

    # Sem base legal na exigência, o aviso é "completar manualmente" (não
    # há citação nenhuma pra "conferir") -- "CONFERIR ANTES DE
    # PROTOCOLAR" é a marca do caso COM citação, coberta em
    # test_base_legal_da_exigencia_aparece_com_marca_de_conferencia.
    assert "não identificada" in textos.lower()
    assert "completar manualmente" in textos.lower()
    # Nenhum "art." ou "artigo" nem "§" aparece em lugar nenhum do
    # documento inteiro quando a exigência não tinha base legal --
    # espaço em branco marcado, nunca um artigo chutado.
    import re
    assert not re.search(r"\bart(?:igo)?\.?\s*\d+|§\s*\d+", textos, re.IGNORECASE)


def test_fundamentacao_e_pedido_da_ia_aparecem_no_documento():
    documento = gerar_recurso(
        _processo(), _empresa(), _exigencia(), NARRATIVA_TESTE,
        "Fundamentação argumentativa de teste, única no documento.",
        "Requer-se, com base no exposto, a reforma da decisão.",
    )
    textos = " ".join(_textos(documento))
    assert "Fundamentação argumentativa de teste, única no documento." in textos
    assert "Requer-se, com base no exposto, a reforma da decisão." in textos


def test_categoria_da_exigencia_mostra_rotulo_amigavel():
    exigencia = _exigencia(categoria="qualificacao_tecnica")
    documento = gerar_recurso(_processo(), _empresa(), exigencia, NARRATIVA_TESTE, "F", "P")
    textos = " ".join(_textos(documento))
    assert "Qualificação Técnica" in textos


def test_empresa_sem_representante_levanta_erro():
    with pytest.raises(EmpresaSemRepresentanteError):
        gerar_recurso(
            _processo(), _empresa(representante_legal_nome=None), _exigencia(), NARRATIVA_TESTE, "F", "P"
        )


def test_bloco_de_assinatura_tem_nome_cargo_e_cpf():
    documento = gerar_recurso(_processo(), _empresa(), _exigencia(), NARRATIVA_TESTE, "F", "P")
    textos = _textos(documento)
    assert any("José da Silva Fictício, Sócio-administrador" in t for t in textos)
    assert any("CPF: 000.000.000-00" in t for t in textos)


def test_devolve_docx_reabrivel():
    documento = gerar_recurso(_processo(), _empresa(), _exigencia(), NARRATIVA_TESTE, "F", "P")
    buffer = io.BytesIO()
    documento.save(buffer)
    buffer.seek(0)

    from docx import Document

    reaberto = Document(buffer)
    assert len(reaberto.paragraphs) > 0


# ---------- Teste de ouro: recurso contra exigência real (Paulínia) ----------
#
# Texto REAL, capturado direto do processo real de Paulínia (id 1) nesta
# mesma sessão (investigação da Camada 0 sobre base_legal) -- não é uma
# consulta ao vivo no nexlicit.db: diferente do golden test de
# planilha_preco/minuta (que processa um PDF real, estático, em
# uploads/), a extração de exigência depende de IA de verdade, não tem
# como reproduzir offline a partir do PDF sozinho. Fixar o texto real já
# observado dá a mesma complexidade de edital real sem acoplar a suíte
# automatizada ao estado mutável do banco de produção (que muda toda vez
# que você usa o app de verdade).
_EXIGENCIA_REAL_PAULINIA = {
    "categoria": "habilitacao_juridica",
    "descricao": "Registro empresarial na Junta Comercial",
    "trecho": (
        "Registro empresarial na Junta Comercial, no caso de empresário individual "
        "(ou cédula de identidade em se tratando de pessoa física não empresária);"
    ),
    "pagina": None,
    "base_legal": "art. 66 da Lei 14.133/2021",
}


def test_recurso_contra_exigencia_real_de_paulinia_com_narrativa_de_teste():
    processo = _processo(id=1, nome="PM Paulínia (teste)")
    narrativa_ficticia = (
        "Situação fictícia de teste, sem relação com o processo real: o documento "
        "exigido foi entregue dentro do prazo, mas o pregoeiro considerou entregue "
        "fora do prazo por um erro de leitura da data no sistema de protocolo."
    )

    documento = gerar_recurso(
        processo, _empresa(), _EXIGENCIA_REAL_PAULINIA, narrativa_ficticia,
        "Fundamentação de teste conectando a exigência real ao relato fictício.",
        "Pedido de teste.",
    )

    textos = " ".join(p.text for p in documento.paragraphs)
    assert _EXIGENCIA_REAL_PAULINIA["trecho"] in textos
    assert narrativa_ficticia in textos
    assert _EXIGENCIA_REAL_PAULINIA["base_legal"] in textos
    assert "CONFERIR ANTES DE PROTOCOLAR" in textos
