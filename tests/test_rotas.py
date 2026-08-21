# Testes das rotas (app/rotas/, plugadas em app/main.py) via TestClient.
#
# Isolamento do banco/uploads: a fixture cliente_teste troca DATABASE_PATH e
# UPLOAD_DIR pra caminhos dentro do tmp_path do pytest, via monkeypatch nos
# módulos que já importaram essas constantes pro próprio namespace (mesmo
# padrão usado pra isolar GEMINI_API_KEY no Passo 3) — não dá pra só trocar
# app.config.DATABASE_PATH, isso não afeta quem já importou o valor antes.
#
# Nenhum teste aqui chama a API real do Gemini: o único teste que passa pela
# análise (que chamaria extrair_checklist) faz monkeypatch nela também.

import fitz
import pytest
from fastapi.testclient import TestClient


def _criar_pdf_exemplo(caminho) -> None:
    documento = fitz.open()
    pagina = documento.new_page()
    pagina.insert_text(
        (72, 72), "Clausula 5. O licitante deve apresentar Certidao Negativa de Debitos."
    )
    documento.save(str(caminho))
    documento.close()


def _checklist_falso(texto_completo, contexto_processo):
    assert "Certidao Negativa de Debitos" in texto_completo
    assert isinstance(contexto_processo, dict)
    return [
        {
            "categoria": "habilitacao_fiscal_social_trabalhista",
            "descricao": "Certidão Negativa de Débitos",
            "base_legal": None,
            "trecho": "O licitante deve apresentar Certidao Negativa de Debitos.",
            "obrigatorio_para": "todos",
        }
    ]


@pytest.fixture
def cliente_teste(tmp_path, monkeypatch) -> TestClient:
    caminho_db = str(tmp_path / "teste.db")
    pasta_uploads = tmp_path / "uploads"
    pasta_uploads.mkdir()

    monkeypatch.setattr("app.db.conexao.DATABASE_PATH", caminho_db)
    monkeypatch.setattr("app.rotas.processos.UPLOAD_DIR", str(pasta_uploads))

    from app.main import app

    return TestClient(app)


def test_criar_processo_com_upload(cliente_teste, tmp_path):
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)

    with open(caminho_pdf, "rb") as arquivo:
        resposta = cliente_teste.post(
            "/processos",
            data={"nome": "Dispensa 001/2026", "orgao": "Prefeitura Teste"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )

    assert resposta.status_code == 201
    assert "id" in resposta.json()


def test_listar_processos(cliente_teste):
    cliente_teste.post("/processos", data={"nome": "Processo A"}, files={})
    cliente_teste.post("/processos", data={"nome": "Processo B"}, files={})

    resposta = cliente_teste.get("/processos")

    assert resposta.status_code == 200
    nomes = [p["nome"] for p in resposta.json()]
    assert "Processo A" in nomes
    assert "Processo B" in nomes


def test_obter_processo_especifico(cliente_teste):
    resposta_criacao = cliente_teste.post(
        "/processos", data={"nome": "Processo específico"}, files={}
    )
    processo_id = resposta_criacao.json()["id"]

    resposta = cliente_teste.get(f"/processos/{processo_id}")

    assert resposta.status_code == 200
    corpo = resposta.json()
    assert corpo["nome"] == "Processo específico"
    assert corpo["arquivos"] == []
    assert corpo["exigencias"] == []


def test_obter_processo_inexistente_devolve_404(cliente_teste):
    resposta = cliente_teste.get("/processos/999999")

    assert resposta.status_code == 404
    assert "999999" in resposta.json()["detail"]


def test_analisar_processo_com_formato_nao_suportado_devolve_400(cliente_teste, tmp_path):
    caminho_invalido = tmp_path / "arquivo.xyz"
    caminho_invalido.write_text("conteudo qualquer")

    with open(caminho_invalido, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Processo com arquivo ruim"},
            files={"arquivos": ("arquivo.xyz", arquivo, "application/octet-stream")},
        )
    processo_id = resposta_criacao.json()["id"]

    resposta = cliente_teste.post(f"/processos/{processo_id}/analisar")

    assert resposta.status_code == 400


def test_analisar_processo_sem_nenhum_arquivo_enviado_devolve_400(cliente_teste):
    resposta_criacao = cliente_teste.post(
        "/processos", data={"nome": "Processo sem upload nenhum"}, files={}
    )
    processo_id = resposta_criacao.json()["id"]

    resposta = cliente_teste.post(f"/processos/{processo_id}/analisar")

    assert resposta.status_code == 400
    assert "arquivos" in resposta.json()["detail"]


def test_fluxo_completo_analisar_e_atualizar_status(cliente_teste, tmp_path, monkeypatch):
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)

    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Processo completo"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]

    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_falso)

    resposta_analise = cliente_teste.post(f"/processos/{processo_id}/analisar")
    assert resposta_analise.status_code == 200
    resumo = resposta_analise.json()
    assert resumo["total_exigencias"] == 1
    assert resumo["localizadas"] == 1

    # Reprocessar sem forçar deve recusar com 409.
    resposta_repetida = cliente_teste.post(f"/processos/{processo_id}/analisar")
    assert resposta_repetida.status_code == 409

    resposta_processo = cliente_teste.get(f"/processos/{processo_id}")
    exigencias = resposta_processo.json()["exigencias"]
    assert len(exigencias) == 1
    id_exigencia = exigencias[0]["id"]

    resposta_patch = cliente_teste.patch(
        f"/exigencias/{id_exigencia}",
        json={"status_check": "ok", "observacao": "revisado no teste"},
    )
    assert resposta_patch.status_code == 200
    corpo_patch = resposta_patch.json()
    assert corpo_patch["status_check"] == "ok"
    assert corpo_patch["observacao_usuario"] == "revisado no teste"


def test_reprocessar_com_forcar_substitui_dados_antigos(cliente_teste, tmp_path, monkeypatch):
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)

    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Processo para reprocessar"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]

    def checklist_primeira_analise(texto_completo, contexto_processo):
        assert "Certidao Negativa de Debitos" in texto_completo
        assert isinstance(contexto_processo, dict)
        return [
            {
                "categoria": "habilitacao_fiscal_social_trabalhista",
                "descricao": "Exigência da primeira análise",
                "base_legal": None,
                "trecho": "O licitante deve apresentar Certidao Negativa de Debitos.",
                "obrigatorio_para": "todos",
            }
        ]

    # 1ª análise: funciona normalmente.
    monkeypatch.setattr("app.pipeline.extrair_checklist", checklist_primeira_analise)
    resposta_1 = cliente_teste.post(f"/processos/{processo_id}/analisar")
    assert resposta_1.status_code == 200
    assert resposta_1.json()["total_exigencias"] == 1

    exigencias_apos_1a = cliente_teste.get(f"/processos/{processo_id}").json()["exigencias"]
    assert len(exigencias_apos_1a) == 1
    assert exigencias_apos_1a[0]["descricao"] == "Exigência da primeira análise"

    # 2ª análise sem forçar: recusa com 409, nada muda.
    resposta_2_sem_forcar = cliente_teste.post(f"/processos/{processo_id}/analisar")
    assert resposta_2_sem_forcar.status_code == 409

    def checklist_segunda_analise(texto_completo, contexto_processo):
        assert "Certidao Negativa de Debitos" in texto_completo
        assert isinstance(contexto_processo, dict)
        return [
            {
                "categoria": "declaracoes_exigidas",
                "descricao": "Exigência da segunda análise (forçada)",
                "base_legal": None,
                "trecho": "O licitante deve apresentar Certidao Negativa de Debitos.",
                "obrigatorio_para": "todos",
            },
            {
                "categoria": "requisitos_proposta",
                "descricao": "Outra exigência da segunda análise",
                "base_legal": None,
                "trecho": "não existe esse trecho no documento de exemplo",
                "obrigatorio_para": "vencedor",
            },
        ]

    # 2ª análise com ?forcar=true: reprocessa de verdade.
    monkeypatch.setattr("app.pipeline.extrair_checklist", checklist_segunda_analise)
    resposta_2_com_forcar = cliente_teste.post(f"/processos/{processo_id}/analisar?forcar=true")
    assert resposta_2_com_forcar.status_code == 200
    assert resposta_2_com_forcar.json()["total_exigencias"] == 2

    exigencias_apos_2a = cliente_teste.get(f"/processos/{processo_id}").json()["exigencias"]
    descricoes = {e["descricao"] for e in exigencias_apos_2a}

    # Os dados antigos sumiram (não ficaram 3 exigências, viraram 2 novas).
    assert len(exigencias_apos_2a) == 2
    assert "Exigência da primeira análise" not in descricoes
    assert descricoes == {
        "Exigência da segunda análise (forçada)",
        "Outra exigência da segunda análise",
    }


def test_atualizar_status_de_exigencia_inexistente_devolve_404(cliente_teste):
    resposta = cliente_teste.patch("/exigencias/999999", json={"status_check": "ok"})

    assert resposta.status_code == 404


def test_analisar_processo_sem_gemini_api_key_devolve_500(cliente_teste, tmp_path, monkeypatch):
    # Arquivo bom (passa pela extração) — o problema tem que ser só a
    # configuração ausente, não o formato do arquivo enviado. Não faz
    # monkeypatch de extrair_checklist: queremos o caminho real até
    # _chamar_gemini, que barra ANTES de qualquer chamada de rede.
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)

    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Processo sem chave configurada"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]

    monkeypatch.setattr("app.ia.llm_client.GEMINI_API_KEY", "")

    resposta = cliente_teste.post(f"/processos/{processo_id}/analisar")

    assert resposta.status_code == 500
    detalhe = resposta.json()["detail"]
    assert "GEMINI_API_KEY" in detalhe
    assert "configuração do servidor" in detalhe


# ---------- POST /processos/{id}/perguntar (Fase 2, Camada 1) ----------


def test_perguntar_devolve_resposta_da_ia(cliente_teste, tmp_path, monkeypatch):
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)

    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Processo com pergunta"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]

    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_falso)
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    def ia_falsa(texto_completo_marcado, pergunta):
        assert texto_completo_marcado.startswith("[PÁGINA 1]")
        assert pergunta == "Qual certidão é exigida?"
        return {"encontrado": True, "resposta": "Certidão Negativa de Débitos.", "paginas": [1]}

    monkeypatch.setattr("app.pipeline._responder_pergunta_ia", ia_falsa)

    resposta = cliente_teste.post(
        f"/processos/{processo_id}/perguntar", json={"pergunta": "Qual certidão é exigida?"}
    )

    assert resposta.status_code == 200
    corpo = resposta.json()
    assert corpo["encontrado"] is True
    assert corpo["paginas"] == [1]


def test_perguntar_processo_inexistente_devolve_404(cliente_teste):
    resposta = cliente_teste.post("/processos/999/perguntar", json={"pergunta": "qualquer"})
    assert resposta.status_code == 404


def test_perguntar_processo_sem_analise_devolve_400(cliente_teste, tmp_path):
    with open(tmp_path / "vazio.pdf", "wb"):
        pass  # não precisa de conteúdo, só cria o processo sem analisar

    resposta_criacao = cliente_teste.post("/processos", data={"nome": "Processo não analisado"})
    processo_id = resposta_criacao.json()["id"]

    resposta = cliente_teste.post(
        f"/processos/{processo_id}/perguntar", json={"pergunta": "qualquer"}
    )

    assert resposta.status_code == 400
    assert str(processo_id) in resposta.json()["detail"]


# ---------- POST /processos/{id}/detectar-inconsistencias (motor de inconsistências, Camada 1) ----------


def _criar_pdf_com_tr_embutido(caminho) -> None:
    # 2 páginas: a primeira só cita o TR de passagem (não é o título — a
    # Camada 0 real precisa saber ignorar isso), a segunda tem o título
    # isolado de verdade, cada rótulo na sua própria linha (mesmo formato
    # calibrado contra os editais reais em app/inconsistencias/limite_tr.py).
    documento = fitz.open()
    pagina1 = documento.new_page()
    pagina1.insert_text(
        (72, 72), "Clausula 5. Conforme disposto no Termo de Referencia, anexo a este edital."
    )
    pagina2 = documento.new_page()
    pagina2.insert_text((72, 72), "ANEXO I")
    pagina2.insert_text((72, 100), "TERMO DE REFERENCIA")
    pagina2.insert_text((72, 130), "1. OBJETO")
    pagina2.insert_text((72, 160), "Fornecimento de bens diversos.")
    documento.save(str(caminho))
    documento.close()


def test_detectar_inconsistencias_devolve_achados_da_ia(cliente_teste, tmp_path, monkeypatch):
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_com_tr_embutido(caminho_pdf)

    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Processo com inconsistência"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]

    # Fake local (não o _checklist_falso compartilhado, que exige o texto
    # do PDF genérico de outros testes) — Camada 1 de inconsistências não
    # depende do checklist, o conteúdo devolvido aqui é irrelevante.
    monkeypatch.setattr("app.pipeline.extrair_checklist", lambda texto, contexto: [])
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    def ia_falsa(texto_edital_marcado, texto_tr_marcado):
        assert texto_edital_marcado.startswith("[PÁGINA 1]")
        assert texto_tr_marcado.startswith("[PÁGINA 2]")
        return [
            {
                "tipo": "prazo",
                "descricao": "prazo diverge",
                "trecho_edital": "trecho do edital",
                "pagina_edital": 1,
                "trecho_tr": "trecho do TR",
                "pagina_tr": 2,
            }
        ]

    monkeypatch.setattr("app.pipeline._detectar_inconsistencias_ia", ia_falsa)

    resposta = cliente_teste.post(f"/processos/{processo_id}/detectar-inconsistencias")

    assert resposta.status_code == 200
    corpo = resposta.json()
    assert corpo["comparacao_possivel"] is True
    assert corpo["motivo_impossibilidade"] is None
    assert len(corpo["inconsistencias"]) == 1
    assert corpo["inconsistencias"][0]["tipo"] == "prazo"


def test_detectar_inconsistencias_tr_nao_identificado_devolve_200_com_motivo(
    cliente_teste, tmp_path, monkeypatch
):
    # PDF de 1 página só, sem nenhum marcador de TR — a Camada 0 real não
    # identifica nada. Isso NÃO é um erro HTTP: é uma resposta válida do
    # produto (mesmo princípio de "encontrado": false no Q&A).
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)

    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Processo sem TR identificável"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]

    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_falso)
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    resposta = cliente_teste.post(f"/processos/{processo_id}/detectar-inconsistencias")

    assert resposta.status_code == 200
    corpo = resposta.json()
    assert corpo["comparacao_possivel"] is False
    assert corpo["motivo_impossibilidade"] is not None
    assert corpo["inconsistencias"] == []


def test_detectar_inconsistencias_processo_inexistente_devolve_404(cliente_teste):
    resposta = cliente_teste.post("/processos/999/detectar-inconsistencias")
    assert resposta.status_code == 404


# ---------- Modo estático do demo público (Fase 3, app/demo_estatico.py) ----------
#
# DEMO_ESTATICO é lido do ambiente uma vez, na importação de app/config.py
# (mesmo padrão de GEMINI_API_KEY) — por isso o monkeypatch aqui é direto
# em app.demo_estatico.DEMO_ESTATICO (o nome que a função de bloqueio
# consulta), não na variável de ambiente, que não teria efeito depois do
# módulo já importado.


def test_criar_processo_bloqueado_em_modo_demo_estatico(cliente_teste, monkeypatch):
    monkeypatch.setattr("app.demo_estatico.DEMO_ESTATICO", True)

    resposta = cliente_teste.post("/processos", data={"nome": "Não deveria criar"})

    assert resposta.status_code == 403
    assert "demonstração" in resposta.json()["detail"]


def test_analisar_processo_bloqueado_em_modo_demo_estatico(cliente_teste, monkeypatch):
    # Processo criado ANTES do modo estático ligar (senão a própria criação
    # já bloquearia) — confirma que /analisar bloqueia mesmo pra um
    # processo que já existe de verdade, não só quando o id é inventado.
    resposta_criacao = cliente_teste.post("/processos", data={"nome": "Processo existente"})
    processo_id = resposta_criacao.json()["id"]

    monkeypatch.setattr("app.demo_estatico.DEMO_ESTATICO", True)
    resposta = cliente_teste.post(f"/processos/{processo_id}/analisar")

    assert resposta.status_code == 403


def test_perguntar_bloqueado_em_modo_demo_estatico(cliente_teste, monkeypatch):
    monkeypatch.setattr("app.demo_estatico.DEMO_ESTATICO", True)

    resposta = cliente_teste.post("/processos/999/perguntar", json={"pergunta": "qualquer"})

    assert resposta.status_code == 403


def test_detectar_inconsistencias_bloqueado_em_modo_demo_estatico(cliente_teste, monkeypatch):
    monkeypatch.setattr("app.demo_estatico.DEMO_ESTATICO", True)

    resposta = cliente_teste.post("/processos/999/detectar-inconsistencias")

    assert resposta.status_code == 403


def test_listar_e_obter_processo_continuam_liberados_em_modo_demo_estatico(
    cliente_teste, monkeypatch
):
    # Modo estático bloqueia só escrita — leitura (listar, obter) continua
    # liberada, porque é exatamente o que o demo público precisa mostrar.
    resposta_criacao = cliente_teste.post("/processos", data={"nome": "Antes do modo estático"})
    processo_id = resposta_criacao.json()["id"]

    monkeypatch.setattr("app.demo_estatico.DEMO_ESTATICO", True)

    assert cliente_teste.get("/processos").status_code == 200
    assert cliente_teste.get(f"/processos/{processo_id}").status_code == 200


# ---------- POST /processos/{id}/gerar-declaracoes (Fase 4, Camada 1) ----------


def _dados_empresa_teste() -> dict:
    return {
        "razao_social": "Exemplo Fornecedora de Materiais Ltda",
        "cnpj": "12.345.678/0001-90",
        "representante_legal_nome": "José da Silva Fictício",
        "representante_legal_cargo": "Sócio-administrador",
        "representante_legal_cpf": "000.000.000-00",
    }


def _checklist_com_declaracao(texto_completo, contexto_processo):
    return [
        {
            "categoria": "declaracoes_exigidas",
            "descricao": "Declaração de não emprego de menor",
            "base_legal": None,
            "trecho": "O licitante deve apresentar Certidao Negativa de Debitos.",
            "obrigatorio_para": "todos",
        }
    ]


def test_gerar_declaracoes_devolve_docx_valido(cliente_teste, tmp_path, monkeypatch):
    from app.db.repositorio import criar_empresa

    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)
    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Pregão Eletrônico 01/2026"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]

    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_com_declaracao)
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    empresa_id = criar_empresa(_dados_empresa_teste())

    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-declaracoes?empresa_id={empresa_id}")

    assert resposta.status_code == 200
    assert resposta.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in resposta.headers["content-disposition"]

    # O corpo da resposta precisa ser um .docx de verdade, reabrível --
    # não só bytes com o content-type certo.
    import io as _io

    import docx as _docx

    documento = _docx.Document(_io.BytesIO(resposta.content))
    texto_completo = " ".join(p.text for p in documento.paragraphs)
    assert "Exemplo Fornecedora de Materiais Ltda" in texto_completo
    assert "Certidao Negativa de Debitos" in texto_completo
    assert "José da Silva Fictício" in texto_completo


def test_gerar_declaracoes_processo_inexistente_devolve_404(cliente_teste):
    from app.db.repositorio import criar_empresa

    empresa_id = criar_empresa(_dados_empresa_teste())

    resposta = cliente_teste.post(f"/processos/999999/gerar-declaracoes?empresa_id={empresa_id}")

    assert resposta.status_code == 404


def test_gerar_declaracoes_empresa_inexistente_devolve_404(cliente_teste, tmp_path, monkeypatch):
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)
    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos", data={"nome": "Processo teste"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]
    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_com_declaracao)
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-declaracoes?empresa_id=999999")

    assert resposta.status_code == 404


def test_gerar_declaracoes_processo_sem_declaracoes_devolve_400(cliente_teste, tmp_path, monkeypatch):
    from app.db.repositorio import criar_empresa

    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)
    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos", data={"nome": "Processo teste"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]
    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_falso)  # sem categoria declaracoes_exigidas
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    empresa_id = criar_empresa(_dados_empresa_teste())
    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-declaracoes?empresa_id={empresa_id}")

    assert resposta.status_code == 400
    assert "declaracoes_exigidas" in resposta.json()["detail"] or "Declarações" in resposta.json()["detail"]


def test_gerar_declaracoes_empresa_sem_representante_devolve_400(cliente_teste, tmp_path, monkeypatch):
    from app.db.repositorio import criar_empresa

    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)
    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos", data={"nome": "Processo teste"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]
    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_com_declaracao)
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    dados_sem_representante = {**_dados_empresa_teste(), "representante_legal_nome": None}
    empresa_id = criar_empresa(dados_sem_representante)

    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-declaracoes?empresa_id={empresa_id}")

    assert resposta.status_code == 400



# ---------- Planilha de preço (Fase 4, Camada 1, decisão B — 16/08/2026) ----------


def _criar_pdf_com_tabela_itens(caminho) -> None:
    documento = fitz.open()
    pagina = documento.new_page()
    pagina.insert_text(
        (72, 72),
        "ITEM DESCRICAO DO PRODUTO UNIDADE QUANTIDADE\n"
        "1 TERMOMETRO DIGITAL. UND 10\n"
        "2 SERINGA DESCARTAVEL. UND 50\n",
    )
    documento.save(str(caminho))
    documento.close()


def _checklist_falso_generico(texto_completo, contexto_processo):
    return [
        {
            "categoria": "habilitacao_fiscal_social_trabalhista",
            "descricao": "Certidão Negativa de Débitos",
            "base_legal": None,
            "trecho": texto_completo[:40],
            "obrigatorio_para": "todos",
        }
    ]


def _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch) -> int:
    caminho_pdf = tmp_path / "edital_com_tabela.pdf"
    _criar_pdf_com_tabela_itens(caminho_pdf)
    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos",
            data={"nome": "Pregão com tabela de itens"},
            files={"arquivos": ("edital_com_tabela.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]

    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_falso_generico)
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    return processo_id


def test_salvar_preco_item_via_patch_persiste(cliente_teste, tmp_path, monkeypatch):
    processo_id = _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch)

    resposta = cliente_teste.patch(
        f"/processos/{processo_id}/itens/1/preco",
        json={"quantidade": 10, "preco_unitario": 25.5},
    )

    assert resposta.status_code == 200
    corpo = resposta.json()
    assert corpo["numero_item"] == 1
    assert corpo["quantidade"] == 10
    assert corpo["preco_unitario"] == 25.5

    from app.db.repositorio import obter_precos_item

    precos = obter_precos_item(processo_id)
    assert precos[1]["preco_unitario"] == 25.5


def test_salvar_preco_item_via_patch_atualiza_sem_duplicar(cliente_teste, tmp_path, monkeypatch):
    processo_id = _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch)

    cliente_teste.patch(f"/processos/{processo_id}/itens/1/preco", json={"quantidade": 10, "preco_unitario": 25.5})
    resposta = cliente_teste.patch(f"/processos/{processo_id}/itens/1/preco", json={"quantidade": 12, "preco_unitario": 30.0})

    assert resposta.status_code == 200

    from app.db.repositorio import obter_precos_item

    precos = obter_precos_item(processo_id)
    assert len(precos) == 1
    assert precos[1]["quantidade"] == 12


def test_gerar_planilha_preco_devolve_xlsx_valido(cliente_teste, tmp_path, monkeypatch):
    processo_id = _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch)
    cliente_teste.patch(f"/processos/{processo_id}/itens/1/preco", json={"quantidade": 10, "preco_unitario": 25.5})
    cliente_teste.patch(f"/processos/{processo_id}/itens/2/preco", json={"quantidade": 20, "preco_unitario": 3.0})

    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-planilha-preco")

    assert resposta.status_code == 200
    assert resposta.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    assert "attachment" in resposta.headers["content-disposition"]

    # Corpo precisa ser um XLSX de verdade, reabrível — não só bytes com o
    # content-type certo, mesmo princípio já usado pro teste do DOCX.
    import io as _io

    from openpyxl import load_workbook

    planilha_lida = load_workbook(_io.BytesIO(resposta.content)).active
    assert planilha_lida is not None
    assert planilha_lida["A2"].value == 1
    assert planilha_lida["C2"].value == 10
    assert planilha_lida["D2"].value == 25.5
    assert planilha_lida["E2"].value == 255.0


def test_gerar_planilha_preco_processo_inexistente_devolve_404(cliente_teste):
    resposta = cliente_teste.post("/processos/999999/gerar-planilha-preco")
    assert resposta.status_code == 404


def test_gerar_planilha_preco_processo_sem_catalogo_devolve_400(cliente_teste, tmp_path, monkeypatch):
    # Processo analisado normalmente (PDF sem tabela de itens reconhecível)
    # — catálogo fica vazio, gerar a planilha não trava o servidor, devolve
    # erro claro em vez de um XLSX vazio sem sentido.
    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)
    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos", data={"nome": "Processo sem tabela"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]
    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_falso_generico)
    cliente_teste.post(f"/processos/{processo_id}/analisar")

    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-planilha-preco")

    assert resposta.status_code == 400


# ---------- Minuta de proposta (Fase 4, Camada 1, 19/08/2026) ----------


def test_salvar_marca_fabricante_modelo_via_patch_persiste(cliente_teste, tmp_path, monkeypatch):
    processo_id = _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch)

    resposta = cliente_teste.patch(
        f"/processos/{processo_id}/itens/1/preco",
        json={"quantidade": 5, "preco_unitario": 200.0, "marca": "MarcaX", "fabricante": "FabricanteY", "modelo": "ModeloZ"},
    )

    assert resposta.status_code == 200
    corpo = resposta.json()
    assert corpo["marca"] == "MarcaX"
    assert corpo["fabricante"] == "FabricanteY"
    assert corpo["modelo"] == "ModeloZ"


def test_salvar_validade_proposta_via_patch_persiste(cliente_teste, tmp_path, monkeypatch):
    processo_id = _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch)

    resposta = cliente_teste.patch(
        f"/processos/{processo_id}/validade-proposta", json={"validade_proposta": "60 dias"}
    )

    assert resposta.status_code == 200
    assert resposta.json()["validade_proposta"] == "60 dias"

    from app.db.repositorio import obter_processo

    processo = obter_processo(processo_id)
    assert processo is not None
    assert processo["validade_proposta"] == "60 dias"


def test_salvar_validade_proposta_processo_inexistente_devolve_404(cliente_teste):
    resposta = cliente_teste.patch("/processos/999999/validade-proposta", json={"validade_proposta": "60 dias"})
    assert resposta.status_code == 404


def test_gerar_minuta_devolve_docx_valido(cliente_teste, tmp_path, monkeypatch):
    from app.db.repositorio import criar_empresa

    processo_id = _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch)
    cliente_teste.patch(
        f"/processos/{processo_id}/itens/1/preco",
        json={"quantidade": 10, "preco_unitario": 25.5, "marca": "MarcaX"},
    )
    cliente_teste.patch(f"/processos/{processo_id}/validade-proposta", json={"validade_proposta": "60 dias"})
    empresa_id = criar_empresa(_dados_empresa_teste())

    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-minuta?empresa_id={empresa_id}")

    assert resposta.status_code == 200
    assert resposta.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in resposta.headers["content-disposition"]

    import io as _io

    import docx as _docx

    documento = _docx.Document(_io.BytesIO(resposta.content))
    texto_completo = " ".join(p.text for p in documento.paragraphs)
    assert "Exemplo Fornecedora de Materiais Ltda" in texto_completo
    assert "60 dias" in texto_completo
    assert len(documento.tables) == 1
    assert documento.tables[0].rows[1].cells[2].text.strip() != ""  # marca preenchida foi pra tabela


def test_gerar_minuta_processo_inexistente_devolve_404(cliente_teste):
    from app.db.repositorio import criar_empresa

    empresa_id = criar_empresa(_dados_empresa_teste())
    resposta = cliente_teste.post(f"/processos/999999/gerar-minuta?empresa_id={empresa_id}")
    assert resposta.status_code == 404


def test_gerar_minuta_empresa_inexistente_devolve_404(cliente_teste, tmp_path, monkeypatch):
    processo_id = _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch)
    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-minuta?empresa_id=999999")
    assert resposta.status_code == 404


def test_gerar_minuta_processo_sem_catalogo_devolve_400(cliente_teste, tmp_path, monkeypatch):
    from app.db.repositorio import criar_empresa

    caminho_pdf = tmp_path / "edital.pdf"
    _criar_pdf_exemplo(caminho_pdf)
    with open(caminho_pdf, "rb") as arquivo:
        resposta_criacao = cliente_teste.post(
            "/processos", data={"nome": "Processo sem tabela"},
            files={"arquivos": ("edital.pdf", arquivo, "application/pdf")},
        )
    processo_id = resposta_criacao.json()["id"]
    monkeypatch.setattr("app.pipeline.extrair_checklist", _checklist_falso_generico)
    cliente_teste.post(f"/processos/{processo_id}/analisar")
    empresa_id = criar_empresa(_dados_empresa_teste())

    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-minuta?empresa_id={empresa_id}")

    assert resposta.status_code == 400


def test_gerar_minuta_empresa_sem_representante_devolve_400(cliente_teste, tmp_path, monkeypatch):
    from app.db.repositorio import criar_empresa

    processo_id = _criar_processo_com_catalogo(cliente_teste, tmp_path, monkeypatch)
    dados_sem_representante = {**_dados_empresa_teste(), "representante_legal_nome": None}
    empresa_id = criar_empresa(dados_sem_representante)

    resposta = cliente_teste.post(f"/processos/{processo_id}/gerar-minuta?empresa_id={empresa_id}")

    assert resposta.status_code == 400
