# Geração determinística de declarações preenchidas (Fase 4, Camada 1) --
# monta um DOCX a partir das exigências de categoria "declaracoes_exigidas"
# já extraídas e VALIDADAS no checklist (Passo 3/4). Sem chamada de IA
# nova -- mesmo princípio do Passo 8 (requisitos por item): reaproveitar
# dado já confiável, montagem puramente determinística de documento.
#
# Formato do documento: mesma estrutura de "Declaração Unificada" já
# vista nos próprios editais testados nesta sessão (ex.: Anexo II do
# edital fictício de demo, e do edital real de Toledo/PR) -- um
# preâmbulo ÚNICO ("A empresa X, CNPJ Y, DECLARA que:") seguido de itens
# alfabéticos, não "a empresa declara que" repetido em cada item.
from __future__ import annotations

import string
from typing import Any

# Ver comentário equivalente em app/geracao/assinatura.py: "docx.Document"
# é a função fábrica, "docx.document.Document" é a classe de verdade,
# a que serve pra anotar tipo.
from docx import Document as _construir_documento
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.geracao.assinatura import EmpresaSemRepresentanteError, adicionar_bloco_assinatura

# Mesmo --stamp do CSS do app (app/static/css/nexlicit.css) -- cor de
# alerta já estabelecida no projeto, não uma nova só pra isto.
_COR_ALERTA = RGBColor(0x8B, 0x2E, 0x2E)

# Re-exportado por compatibilidade: erros.py e outros módulos importam
# EmpresaSemRepresentanteError direto daqui desde antes do bloco de
# assinatura ser extraído pra app/geracao/assinatura.py (19/08/2026).
__all__ = ["SemDeclaracoesError", "EmpresaSemRepresentanteError", "gerar_declaracoes"]


class SemDeclaracoesError(Exception):
    """O processo não tem nenhuma exigência de categoria
    "declaracoes_exigidas" salva no checklist -- não há o que gerar."""


def _letra(indice: int) -> str:
    """0->a, 1->b, ..., 25->z, 26->aa, 27->ab, ... -- mesma lógica de
    numeração de coluna de planilha, em letras minúsculas. Declarações
    raramente passam de 10-15 itens; a extensão pra "aa" só entraria em
    um caso bem fora do comum, mas não custa cobrir."""
    letras = string.ascii_lowercase
    posicao = indice + 1  # base 1, não 0
    resultado = ""
    while posicao > 0:
        posicao, resto = divmod(posicao - 1, 26)
        resultado = letras[resto] + resultado
    return resultado


def gerar_declaracoes(processo: dict[str, Any], empresa: dict[str, Any]) -> Document:
    """Monta o documento DOCX de declaração unificada pro `processo`,
    assinado pela `empresa`. Espera os dicts no formato que
    app.db.repositorio.obter_processo()/obter_empresa() devolvem.

    Levanta SemDeclaracoesError se o processo não tiver nenhuma exigência
    de categoria "declaracoes_exigidas". Levanta
    EmpresaSemRepresentanteError se a empresa não tiver
    representante_legal_nome cadastrado.

    Item com confianca="inferido" no checklist original recebe uma marca
    visível (texto em negrito vermelho, mesma cor --stamp do resto do
    app) E um comentário nativo do Word ancorado no parágrafo -- as duas
    coisas, não uma ou outra: comentário sozinho é fácil de não notar sem
    abrir o painel de revisão do Word, marca no texto sozinha não
    documenta o motivo. A incerteza que já existia na extração original
    não desaparece só porque virou um documento bonito.
    """
    declaracoes = [e for e in processo["exigencias"] if e["categoria"] == "declaracoes_exigidas"]
    if not declaracoes:
        raise SemDeclaracoesError(
            f"processo {processo['id']} não tem nenhuma exigência de "
            "'Declarações Exigidas' no checklist -- rode a análise antes, "
            "ou confira se o edital realmente pede declaração nessa categoria"
        )
    if not empresa.get("representante_legal_nome"):
        raise EmpresaSemRepresentanteError(
            f"empresa '{empresa['razao_social']}' não tem representante legal "
            "cadastrado -- complete o cadastro antes de gerar o documento"
        )

    documento = _construir_documento()

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("DECLARAÇÃO UNIFICADA")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)

    documento.add_paragraph()

    # Preâmbulo termina em "as seguintes declarações:", não em "...que:" --
    # achado no golden test contra dado real do Paulínia (13/08/2026):
    # "trecho" NEM SEMPRE é uma continuação limpa de "que" (às vezes já
    # vem como "De que não emprega...", às vezes como frase completa tipo
    # "Para microempresas... Declaração de..."). Forçar um "que" na frente
    # de todo item quebrava a gramática ("que De que não emprega...",
    # duplicado) ou simplesmente não fazia sentido pros itens que já são
    # frase própria. Terminar o preâmbulo sem exigir uma continuação
    # específica aceita qualquer formato de trecho com segurança, sem
    # reescrever o texto extraído pra caber num molde que a IA nem sempre
    # segue.
    preambulo = documento.add_paragraph()
    preambulo.add_run(
        f"A empresa {empresa['razao_social']}, inscrita no CNPJ sob nº {empresa['cnpj']}, "
        "por intermédio de seu representante legal, apresenta, para os fins do disposto "
        f"no edital \"{processo['nome']}\", as seguintes declarações:"
    )

    documento.add_paragraph()

    for indice, declaracao in enumerate(declaracoes):
        paragrafo = documento.add_paragraph()
        # Texto do item vai literal (trecho já extraído e validado, sem
        # reescrever) -- só cai pra "descricao" quando não há trecho
        # nenhum (raro; "inferido" às vezes vem sem citação literal).
        texto_base = declaracao.get("trecho") or declaracao["descricao"]
        paragrafo.add_run(f"{_letra(indice)}) {texto_base}")

        if declaracao["confianca"] == "inferido":
            aviso = paragrafo.add_run(
                "  [CONFERIR: extraído como \"inferido\" no checklist original -- "
                "não é citação literal do edital, revisar antes de assinar]"
            )
            aviso.bold = True
            aviso.font.color.rgb = _COR_ALERTA
            documento.add_comment(
                paragrafo.runs,
                text=(
                    "Esta declaração foi extraída com confiança \"inferido\" no "
                    "checklist original -- a IA interpretou o texto do edital em "
                    "vez de copiar uma citação literal. Confira contra o edital "
                    "fonte antes de assinar."
                ),
                author="NexLicit Engine",
            )

    documento.add_paragraph()
    documento.add_paragraph()

    # Bloco de assinatura compartilhado com a minuta de proposta (Fase 4,
    # Camada 1, 19/08/2026) -- ver app/geracao/assinatura.py. A checagem
    # de representante legal já aconteceu no início desta função (linha
    # ~81), antes de montar o corpo inteiro do documento -- a checagem
    # que adicionar_bloco_assinatura() faz de novo aqui nunca dispara na
    # prática, é só defensivo caso algum dia alguém chame esta função sem
    # passar pela checagem cedo.
    adicionar_bloco_assinatura(documento, empresa)

    return documento
