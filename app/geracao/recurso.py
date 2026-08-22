# Geração do recurso administrativo contra inabilitação indevida (Fase 4,
# Camada 1, 19/08/2026) -- documento de maior risco do projeto: diferente
# de declarações/minuta/planilha (montagem 100% determinística), aqui a
# IA escreve argumentação jurídica real (app.pipeline.gerar_recurso_processo
# -> app.ia.llm_client.gerar_argumento_recurso). Este módulo só monta o
# DOCX a partir do que já foi decidido -- nenhuma chamada de IA aqui.
#
# DECISÃO DE SEGURANÇA (confirmada com Bruno, 19/08/2026): a citação de
# base legal é inserida por ESTE módulo, direto do campo "base_legal" já
# validado da exigência -- a IA nunca escreve número de artigo em lugar
# nenhum (ver app.ia.llm_client._SCHEMA_ARGUMENTO_RECURSO). Marcada
# "CONFERIR ANTES DE PROTOCOLAR" SEMPRE, mesmo quando a exigência de
# origem é "localizado" -- mais rígido que declarações/minuta de
# propósito: o risco aqui não é "esse trecho existe no edital" (isso já
# foi conferido), é "essa citação certa foi usada na linha de argumento
# certa pro caso concreto", e só um profissional avaliando o caso
# consegue descartar isso, nunca o sistema sozinho.

from __future__ import annotations

from typing import Any

# Ver comentário equivalente em app/geracao/assinatura.py: "docx.Document"
# é a função fábrica, "docx.document.Document" é a classe de verdade.
from docx import Document as _construir_documento
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.geracao.assinatura import adicionar_bloco_assinatura

# Mesmo --stamp do CSS do app -- cor de alerta já estabelecida no
# projeto (mesma usada em declaracoes.py e planilha_preco.py).
_COR_ALERTA = RGBColor(0x8B, 0x2E, 0x2E)

# Só os 4 rótulos de categoria de habilitação -- não importa
# RÓTULOS_CATEGORIA de app.rotas.paginas de propósito (rotas não deveria
# ser dependência de geracao, mesmo motivo já registrado em
# app.pipeline sobre não importar dali). Duplicar 4 linhas é mais simples
# que reorganizar módulo pra uma dependência cruzada.
_ROTULOS_CATEGORIA_HABILITACAO = {
    "habilitacao_juridica": "Habilitação Jurídica",
    "habilitacao_fiscal_social_trabalhista": "Habilitação Fiscal, Social e Trabalhista",
    "qualificacao_economico_financeira": "Qualificação Econômico-Financeira",
    "qualificacao_tecnica": "Qualificação Técnica",
}


def _adicionar_titulo_de_aviso(documento: Document) -> None:
    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("MINUTA — REVISAR ANTES DE PROTOCOLAR")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    run_titulo.font.color.rgb = _COR_ALERTA

    subtitulo = documento.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.add_run("Recurso administrativo contra inabilitação")


def _adicionar_cabecalho(documento: Document, processo: dict[str, Any], empresa: dict[str, Any]) -> None:
    identificacao = documento.add_paragraph()
    identificacao.add_run(
        f"{empresa['razao_social']}, inscrita no CNPJ sob nº {empresa['cnpj']}, "
        f"por intermédio de seu representante legal, apresenta recurso administrativo "
        f"referente ao processo \"{processo['nome']}\"."
    )
    if processo.get("orgao"):
        linha = documento.add_paragraph()
        linha.add_run("Órgão: ").bold = True
        linha.add_run(processo["orgao"])


def _adicionar_secao(documento: Document, titulo: str) -> None:
    cabecalho = documento.add_paragraph()
    run = cabecalho.add_run(titulo)
    run.bold = True


def _adicionar_fundamento_no_edital(documento: Document, exigencia: dict[str, Any]) -> None:
    _adicionar_secao(documento, "I. Exigência do edital que motivou a inabilitação")

    rotulo_categoria = _ROTULOS_CATEGORIA_HABILITACAO.get(exigencia["categoria"], exigencia["categoria"])
    linha_categoria = documento.add_paragraph()
    linha_categoria.add_run(f"Categoria: {rotulo_categoria}")

    linha_descricao = documento.add_paragraph()
    linha_descricao.add_run(f"Exigência: {exigencia['descricao']}")

    paragrafo_trecho = documento.add_paragraph()
    paragrafo_trecho.add_run("Trecho literal do edital: ").italic = True
    run_trecho = paragrafo_trecho.add_run(f"“{exigencia['trecho']}”")
    run_trecho.italic = True
    if exigencia.get("pagina"):
        paragrafo_trecho.add_run(f" (página {exigencia['pagina']})").italic = True


def _adicionar_fatos_alegados(documento: Document, narrativa: str) -> None:
    _adicionar_secao(documento, "II. Fatos alegados pelo recorrente")

    aviso = documento.add_paragraph()
    run_aviso = aviso.add_run(
        "O relato abaixo é a versão do recorrente sobre o ocorrido -- não verificada pelo sistema."
    )
    run_aviso.italic = True

    # Narrativa vai VERBATIM, nunca reescrita/resumida pelo código nem
    # pela IA -- mesmo princípio de "trecho" no checklist: o que a pessoa
    # escreveu não é parafraseado por ninguém neste processo.
    paragrafo_narrativa = documento.add_paragraph()
    paragrafo_narrativa.add_run(narrativa)


def _adicionar_fundamentacao(documento: Document, fundamentacao: str) -> None:
    _adicionar_secao(documento, "III. Fundamentação")
    documento.add_paragraph().add_run(fundamentacao)


def _adicionar_fundamento_legal(documento: Document, base_legal: str | None) -> None:
    _adicionar_secao(documento, "IV. Fundamento legal")

    paragrafo = documento.add_paragraph()
    if base_legal:
        paragrafo.add_run(f"Base legal identificada nesta exigência do checklist: {base_legal}")
        aviso = paragrafo.add_run("  [CONFERIR ANTES DE PROTOCOLAR]")
        aviso.bold = True
        aviso.font.color.rgb = _COR_ALERTA
        documento.add_comment(
            paragrafo.runs,
            text=(
                "Esta citação de base legal veio da exigência já extraída no checklist, "
                "nunca gerada pela IA na hora de montar este recurso -- mesmo assim, "
                "confira se o dispositivo citado é o correto para a linha de argumento "
                "usada neste caso concreto antes de protocolar. Uma citação correta usada "
                "no lugar errado é um risco que só avaliação profissional do caso resolve."
            ),
            author="NexLicit Engine",
        )
    else:
        aviso = paragrafo.add_run(
            "[Base legal não identificada nesta exigência do checklist -- "
            "completar manualmente antes de protocolar, se aplicável]"
        )
        aviso.bold = True
        aviso.font.color.rgb = _COR_ALERTA


def _adicionar_pedido(documento: Document, pedido: str) -> None:
    _adicionar_secao(documento, "V. Pedido")
    documento.add_paragraph().add_run(pedido)


def gerar_recurso(
    processo: dict[str, Any],
    empresa: dict[str, Any],
    exigencia: dict[str, Any],
    narrativa: str,
    fundamentacao: str,
    pedido: str,
) -> Document:
    """Monta o DOCX do recurso administrativo. Espera `fundamentacao` e
    `pedido` já gerados por app.pipeline.gerar_recurso_processo (que já
    chamou a IA e já confirmou "narrativa_suficiente" -- este módulo não
    valida isso de novo, só monta o documento com o que recebeu).

    Levanta EmpresaSemRepresentanteError (via adicionar_bloco_assinatura)
    se a empresa não tiver representante legal cadastrado.
    """
    documento = _construir_documento()

    _adicionar_titulo_de_aviso(documento)
    documento.add_paragraph()
    _adicionar_cabecalho(documento, processo, empresa)
    documento.add_paragraph()
    _adicionar_fundamento_no_edital(documento, exigencia)
    documento.add_paragraph()
    _adicionar_fatos_alegados(documento, narrativa)
    documento.add_paragraph()
    _adicionar_fundamentacao(documento, fundamentacao)
    documento.add_paragraph()
    _adicionar_fundamento_legal(documento, exigencia.get("base_legal"))
    documento.add_paragraph()
    _adicionar_pedido(documento, pedido)

    documento.add_paragraph()
    documento.add_paragraph()
    adicionar_bloco_assinatura(documento, empresa)

    return documento
