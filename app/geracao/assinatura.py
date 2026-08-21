# Bloco de assinatura compartilhado entre documentos gerados (Fase 4) --
# extraído de app/geracao/declaracoes.py (19/08/2026) quando a minuta de
# proposta (Camada 1 seguinte) precisou exatamente do mesmo bloco: data,
# linha de assinatura, nome + cargo do representante legal, CPF. Mesma
# lógica de app/rotas/nomes_arquivo.py (extraído da mesma forma, mesmo
# motivo) -- evita duplicar em cada módulo de geração novo.

from __future__ import annotations

from typing import Any

# "docx.Document" (import direto) é a FUNÇÃO fábrica que cria o
# documento -- pra anotar TIPO precisa da classe de verdade,
# "docx.document.Document" (mesmo objeto que Document() devolve, só que
# importado do submódulo, não do pacote). Achado ao corrigir um erro
# real do Pyright ("Expected class but received function") -- sem isso,
# "documento: Document" no parâmetro abaixo confundia o checador de tipo.
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class EmpresaSemRepresentanteError(Exception):
    """A empresa selecionada não tem "representante_legal_nome"
    cadastrado -- obrigatório pro bloco de assinatura. Recusa com
    mensagem clara em vez de gerar um documento com lacuna óbvia no
    lugar de quem vai assinar."""


def adicionar_bloco_assinatura(documento: Document, empresa: dict[str, Any]) -> None:
    """Adiciona data, linha de assinatura e identificação do
    representante legal ao final do `documento`. Levanta
    EmpresaSemRepresentanteError se a empresa não tiver
    representante_legal_nome cadastrado -- checar isso é responsabilidade
    de quem chama, ANTES de montar o resto do documento (não faz sentido
    gerar páginas de conteúdo pra descobrir só no final que falta quem
    assina)."""
    if not empresa.get("representante_legal_nome"):
        raise EmpresaSemRepresentanteError(
            f"empresa '{empresa['razao_social']}' não tem representante legal "
            "cadastrado -- complete o cadastro antes de gerar o documento"
        )

    data = documento.add_paragraph()
    data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data.add_run("______________________, ____ de _______________ de ________.")

    documento.add_paragraph()
    documento.add_paragraph()

    assinatura = documento.add_paragraph()
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    assinatura.add_run("____________________________________")

    linha_nome = documento.add_paragraph()
    linha_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
    texto_nome = empresa["representante_legal_nome"]
    if empresa.get("representante_legal_cargo"):
        texto_nome += f", {empresa['representante_legal_cargo']}"
    linha_nome.add_run(texto_nome)

    if empresa.get("representante_legal_cpf"):
        linha_cpf = documento.add_paragraph()
        linha_cpf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        linha_cpf.add_run(f"CPF: {empresa['representante_legal_cpf']}")
