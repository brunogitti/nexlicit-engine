# Geração determinística da minuta de proposta (Fase 4, Camada 1,
# 19/08/2026) -- monta um DOCX com cabeçalho (empresa + processo), tabela
# de itens (catálogo + preço/marca/fabricante/modelo digitados por
# gente), cláusula de validade da proposta e bloco de assinatura. Sem
# chamada de IA nenhuma, mesmo princípio de declaracoes.py e
# planilha_preco.py: montagem pura de documento em cima de dado já
# confiável ou digitado à mão -- nunca inventado.

from __future__ import annotations

from typing import Any

# Ver comentário equivalente em app/geracao/assinatura.py: "docx.Document"
# é a função fábrica, "docx.document.Document" é a classe de verdade.
from docx import Document as _construir_documento
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.geracao.assinatura import adicionar_bloco_assinatura
from app.geracao.planilha_preco import SemCatalogoError

# Mesmo --stamp do CSS do app -- cor de alerta já estabelecida no
# projeto (mesma usada em declaracoes.py pro aviso de "inferido").
_COR_ALERTA = RGBColor(0x8B, 0x2E, 0x2E)

_CABECALHO_TABELA = [
    "Item", "Descrição", "Marca / Fabricante / Modelo",
    "Quantidade", "Preço unitário (R$)", "Preço total (R$)",
]


def _texto_marca_fabricante_modelo(preco: dict[str, Any] | None) -> list[str]:
    """Devolve uma linha de texto por campo preenchido (marca,
    fabricante, modelo) -- só os que existem, nunca inventa os que
    faltam. Lista vazia se não houver preco_item nenhum salvo pra este
    item, ou se os três campos estiverem vazios."""
    if preco is None:
        return []
    linhas = []
    if preco.get("marca"):
        linhas.append(f"Marca: {preco['marca']}")
    if preco.get("fabricante"):
        linhas.append(f"Fabricante: {preco['fabricante']}")
    if preco.get("modelo"):
        linhas.append(f"Modelo: {preco['modelo']}")
    return linhas


def _preencher_celula(celula, linhas: list[str]) -> None:
    """Preenche uma célula de tabela com uma ou mais linhas de texto
    (cada uma seu próprio parágrafo dentro da célula) -- vazia (nenhum
    parágrafo com texto) se `linhas` estiver vazia, não um traço nem
    "N/A" inventado."""
    if not linhas:
        return
    celula.text = linhas[0]
    for linha in linhas[1:]:
        celula.add_paragraph(linha)


def _adicionar_cabecalho_processo(documento: Document, processo: dict[str, Any], empresa: dict[str, Any]) -> None:
    identificacao = documento.add_paragraph()
    identificacao.add_run(
        f"{empresa['razao_social']}, inscrita no CNPJ sob nº {empresa['cnpj']}, "
        f"apresenta a presente proposta comercial referente ao processo \"{processo['nome']}\"."
    )

    if processo.get("orgao"):
        linha = documento.add_paragraph()
        linha.add_run("Órgão: ").bold = True
        linha.add_run(processo["orgao"])
    if processo.get("modalidade"):
        linha = documento.add_paragraph()
        linha.add_run("Modalidade: ").bold = True
        linha.add_run(processo["modalidade"])
    if processo.get("objeto"):
        linha = documento.add_paragraph()
        linha.add_run("Objeto: ").bold = True
        linha.add_run(processo["objeto"])


def _adicionar_tabela_itens(
    documento: Document, catalogo: list[dict[str, Any]], precos: dict[int, dict[str, Any]]
) -> None:
    tabela = documento.add_table(rows=1, cols=len(_CABECALHO_TABELA))
    tabela.style = "Table Grid"
    for celula, texto in zip(tabela.rows[0].cells, _CABECALHO_TABELA):
        celula.text = texto
        celula.paragraphs[0].runs[0].bold = True

    for item in catalogo:
        preco = precos.get(item["numero"])
        quantidade = preco.get("quantidade") if preco else None
        preco_unitario = preco.get("preco_unitario") if preco else None
        preco_total = quantidade * preco_unitario if quantidade is not None and preco_unitario is not None else None

        linha = tabela.add_row()
        linha.cells[0].text = str(item["numero"])
        linha.cells[1].text = item["texto_bruto"]
        _preencher_celula(linha.cells[2], _texto_marca_fabricante_modelo(preco))
        linha.cells[3].text = "" if quantidade is None else _formatar_numero(quantidade)
        linha.cells[4].text = "" if preco_unitario is None else _formatar_numero(preco_unitario)
        linha.cells[5].text = "" if preco_total is None else _formatar_numero(preco_total)


def _formatar_numero(valor: float) -> str:
    # "%g" tira ".0" de número inteiro sem cortar casa decimal de
    # verdade -- mesmo achado (e mesma correção) do template da planilha
    # de preço (app/templates/planilha_preco.html, 16/08/2026): sem isso,
    # todo valor guardado como REAL no SQLite apareceria com ".0" a mais.
    return "%g" % valor


def _adicionar_clausula_validade(documento: Document, validade_proposta: str | None) -> None:
    documento.add_paragraph()
    paragrafo = documento.add_paragraph()
    if validade_proposta:
        paragrafo.add_run(f"O prazo de validade desta proposta é de {validade_proposta}.")
    else:
        # Não inventa prazo nenhum -- marca visível pra completar antes de
        # assinar, mesmo princípio do aviso de "inferido" em
        # declaracoes.py (texto em negrito vermelho).
        aviso = paragrafo.add_run(
            "[PENDENTE: preencher a validade da proposta antes de assinar -- "
            "campo vazio na tela da planilha de preço]"
        )
        aviso.bold = True
        aviso.font.color.rgb = _COR_ALERTA


def gerar_minuta(
    processo: dict[str, Any],
    empresa: dict[str, Any],
    catalogo: list[dict[str, Any]],
    precos: dict[int, dict[str, Any]],
) -> Document:
    """Monta o DOCX da minuta de proposta pro `processo`, assinado pela
    `empresa`, com os itens do `catalogo` precificados conforme `precos`
    (mesmo formato de app.db.repositorio.obter_catalogo_itens()/
    obter_precos_item()).

    Levanta SemCatalogoError se `catalogo` estiver vazio (mesma exceção
    da planilha de preço -- é a mesma condição de fundo: sem catálogo,
    não tem tabela de itens pra montar). Levanta
    EmpresaSemRepresentanteError (via adicionar_bloco_assinatura) se a
    empresa não tiver representante legal cadastrado.

    Item sem marca/fabricante/modelo, sem quantidade ou sem preço: célula
    vazia, nunca um valor inventado -- mesmo princípio já usado na
    planilha de preço (XLSX). Validade da proposta sem preenchimento:
    marca visível em vermelho, não um prazo chutado.
    """
    if not catalogo:
        raise SemCatalogoError(
            f"processo {processo['id']} não tem catálogo de itens salvo -- rode a análise "
            "(ou reprocesse) depois de confirmar que o edital tem uma tabela de itens "
            "reconhecível, ou não há como montar a minuta sem os itens"
        )

    documento = _construir_documento()

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("PROPOSTA COMERCIAL")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)

    documento.add_paragraph()

    _adicionar_cabecalho_processo(documento, processo, empresa)
    documento.add_paragraph()
    _adicionar_tabela_itens(documento, catalogo, precos)
    _adicionar_clausula_validade(documento, processo.get("validade_proposta"))

    documento.add_paragraph()
    documento.add_paragraph()
    adicionar_bloco_assinatura(documento, empresa)

    return documento
